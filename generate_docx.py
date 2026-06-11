from docx import Document
from docx.shared import Pt
import sys

def create_proposal(filepath):
    doc = Document()

    # Title
    doc.add_heading('Proposta Progettuale: Sistema Smart-Agri', level=1)
    
    # Intro
    p = doc.add_paragraph(
        "Il progetto Smart-Agri consiste nello sviluppo di una soluzione software intelligente in grado "
        "di assistere gli agricoltori nel monitoraggio e nell'analisi dello stato di salute delle colture. "
        "Il sistema integra tecnologie di Computer Vision e un'interfaccia conversazionale basata su Dialogflow, "
        "permettendo all'utente di effettuare diagnosi in tempo reale tramite interazioni in linguaggio naturale."
    )

    doc.add_heading('Architettura e Stato di Avanzamento', level=2)
    p = doc.add_paragraph(
        "L'architettura del sistema si fonda su un backend Flask sviluppato in Python che funge da ponte tra l'utente "
        "(tramite Dialogflow) e i moduli di analisi. Le funzionalità principali già implementate e funzionanti "
        "coprono la struttura fondamentale del progetto:\n"
    )
    p.add_run("• Acquisizione Immagini: ").bold = True
    p.add_run("Integrazione con un servizio fotocamera (CameraService) per l'acquisizione di frame video in tempo reale.\n")
    p.add_run("• Computer Vision (VisionService): ").bold = True
    p.add_run("Analisi dell'immagine catturata tramite modelli di apprendimento automatico in grado di identificare "
              "la specie vegetale, valutare lo stato sanitario, stimare la percentuale di area fogliare con anomalie "
              "e conteggiare il numero di piantine rilevate.\n")
    p.add_run("• Assistente Conversazionale: ").bold = True
    p.add_run("Configurazione degli intenti di interazione (es. 'Saluto', 'AnalisiPianta') con dispatching dinamico su architettura webhook (app.py).\n")
    p.add_run("• Persistenza dei Dati (DBRepository): ").bold = True
    p.add_run("Salvataggio automatico delle osservazioni e dei responsi su un database relazionale (MySQL), rispettando "
              "lo schema logico per le entità plants e observations.")

    doc.add_paragraph(
        "Queste funzionalità base già integrate costituiscono oltre il 70% delle specifiche previste. Il restante 30% "
        "comprenderà l'affinamento dei modelli di classificazione, l'aggiunta di intenti per consultare lo storico dei dati salvati "
        "e il miglioramento delle logiche di feedback (suggerimenti all'utente)."
    )

    doc.add_heading('Componenti Modulari', level=3)
    doc.add_paragraph(
        "Il sistema è stato progettato utilizzando pattern architetturali che garantiscono scalabilità e manutenibilità. "
        "La divisione in strati (controller, repository, services) isola le logiche di visione da quelle del database, "
        "consentendo una gestione degli errori centralizzata (Middleware Catch-All) per prevenire crash applicativi "
        "e garantire sempre una risposta di fallback al webhook di Dialogflow."
    )

    doc.add_heading("Scenario d'uso per la demo dell'ultimo incontro", level=2)
    doc.add_paragraph(
        "Per la demo dell’incontro del 18/06/26, verrà presentato il seguente scenario d’uso della soluzione implementata, "
        "in grado di dimostrare l'integrazione di almeno il 70% delle funzionalità previste dal progetto, costituendo l'ossatura "
        "tecnica su cui si baseranno le integrazioni future."
    )

    doc.add_paragraph(
        "Durante la dimostrazione, un utente (simulante un operatore agricolo) avvierà una conversazione con l'agente "
        "virtuale Smart-Agri. Dopo il saluto iniziale, l'utente chiederà al sistema di analizzare una pianta posizionata "
        "di fronte alla webcam (es. formulando l'intento: \"Analizza questa foglia\"). In tempo reale, il backend riceverà "
        "la richiesta via webhook da Dialogflow, attiverà il modulo CameraService per estrarre il frame video e lo passerà "
        "al VisionService. Il modello di Computer Vision processerà l'immagine identificando la specie (es. Pomodoro), "
        "diagnosticando eventuali patologie, calcolando la percentuale di area anomala e contando le piantine. Terminato "
        "il calcolo (con un vincolo di tempo massimo per evitare timeout), il sistema salverà in modo persistente questi dati "
        "all'interno del database MySQL (smart_agri), legandoli allo storico delle osservazioni. L'assistente virtuale "
        "risponderà quindi vocalmente o testualmente all'utente, fornendo il report completo e, qualora la foglia non "
        "fosse stata inquadrata correttamente, suggerendo di riposizionare l'obiettivo. Questo flusso end-to-end dimostra "
        "la solidità dell'infrastruttura di analisi e acquisizione dati, confermando la predisposizione del sistema all'integrazione "
        "del restante 30% delle funzionalità relative all'interrogazione dello storico."
    )

    doc.save(filepath)
    print(f"Document saved to {filepath}")

if __name__ == '__main__':
    create_proposal(sys.argv[1])
